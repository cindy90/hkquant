"""DOCX exporter — renders the investment memo to .docx via python-docx.

Phase 7 MVP: walks the markdown line-by-line and emits headings / bullets /
paragraphs. Tables and rich formatting are intentionally limited; Phase 9
will upgrade if banker workflows demand it.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document

from ...common.schemas import PredictionSnapshot
from ..report_builder import build_memo_markdown


def _add_heading(doc: Any, line: str) -> None:
    """Convert a markdown ATX heading into a python-docx heading."""
    level = 0
    while line.startswith("#"):
        line = line[1:]
        level += 1
    doc.add_heading(line.strip(), level=min(level, 4))


def _is_bullet(line: str) -> bool:
    s = line.lstrip()
    return s.startswith("- ") or s.startswith("* ")


def export_docx(snapshot: PredictionSnapshot) -> bytes:
    """Render the memo to .docx bytes."""
    md = build_memo_markdown(snapshot)
    doc: Any = Document()

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            _add_heading(doc, line)
        elif line.strip() == "---":
            # rule
            doc.add_paragraph("─" * 40)
        elif _is_bullet(line):
            bullet_text = line.lstrip()[2:]
            doc.add_paragraph(bullet_text, style="List Bullet")
        elif line.startswith("|"):
            # Simple table-row passthrough — write as raw paragraph; full
            # table parsing is out of Phase 7 MVP scope.
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ("export_docx",)
